"""
Автоформатирование .docx-отчета по требованиям лабораторной работы.

Что делает:
- формат листа A4 и поля;
- верхний колонтитул: ФИО, группа, отчет по ЛР N;
- нижний колонтитул: сквозная нумерация страниц;
- основной текст: TNR 14, 1.5 интервал, 1.25 см красная строка, по ширине;
- заголовки Heading 1 / Заголовок 1: Arial 14 bold uppercase, нумерация, с новой страницы;
- таблицы: центрирование, запрет разрыва строки таблицы между страницами,
  оформление первой строки и текста ячеек;
- автоматически вставляет/перенумеровывает подписи:
  * перед таблицей: "Таблица N — ..."
  * после рисунка: "Рисунок N — ..."
- умеет нумеровать рисунки и внутри таблиц, не удаляя сами изображения.

Чего код не может гарантировать на 100% автоматически:
- "логичную" ссылку на рисунок/таблицу в предыдущем абзаце;
- запрет окончания строки предлогом/союзом/частицей;
- идеальную подгонку страницы так, чтобы внизу не оставалось больше 1 пустой строки.

Запуск:
1) GUI:
   python lab_report_formatter.py

2) CLI:
   python lab_report_formatter.py input.docx output.docx --fio "Иванов И.И." --group "ИВТ-101" --lab 5

Зависимости:
    pip install python-docx
"""

from __future__ import annotations

import argparse
import os
import re
import tkinter as tk
import webbrowser
from dataclasses import dataclass
from tkinter import filedialog, messagebox
from typing import Iterator, Optional

from docx import Document
from docx.document import Document as _Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

PLACEHOLDER_TITLE = "..."
BLACK = RGBColor(0, 0, 0)
HEADING_STYLE_NAMES = {"Heading 1", "Заголовок 1"}
CAPTION_RE = re.compile(r"^(Рисунок|Таблица)\s+(\d+)\s*[-—]\s*(.*)$", re.IGNORECASE)
LEADING_NUMBER_RE = re.compile(r"^\s*\d+([.)]|\s)+")
SOFT_HYPHEN = "\u00ad"


@dataclass
class Meta:
    fio: str
    group: str
    lab_number: str


def iter_block_items(parent: _Document | _Cell) -> Iterator[Paragraph | Table]:
    """Итерирует параграфы и таблицы в порядке документа."""
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc

    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def iter_paragraphs_recursive(parent: _Document | _Cell) -> Iterator[Paragraph]:
    """Итерирует все параграфы документа, включая вложенные в таблицы."""
    for block in iter_block_items(parent):
        if isinstance(block, Paragraph):
            yield block
        else:
            for row in block.rows:
                for cell in row.cells:
                    yield from iter_paragraphs_recursive(cell)


def iter_blocks_recursive(parent: _Document | _Cell) -> Iterator[Paragraph | Table]:
    """Итерирует блоки документа, включая содержимое таблиц, в естественном порядке."""
    for block in iter_block_items(parent):
        yield block
        if isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    yield from iter_blocks_recursive(cell)


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.xpath(
        ".//*[local-name()='drawing' or local-name()='pict' or local-name()='imagedata' or local-name()='object' or local-name()='OLEObject']"
    ))


def run_has_drawing(run) -> bool:
    return bool(run._element.xpath(
        ".//*[local-name()='drawing' or local-name()='pict' or local-name()='imagedata' or local-name()='object' or local-name()='OLEObject']"
    ))


def paragraph_text_clean(paragraph: Optional[Paragraph]) -> str:
    if paragraph is None:
        return ""
    return (paragraph.text or "").strip()


def is_empty_paragraph(paragraph: Optional[Paragraph]) -> bool:
    return paragraph is not None and not paragraph_text_clean(paragraph) and not paragraph_has_drawing(paragraph)


def is_caption_paragraph(paragraph: Optional[Paragraph], kind: Optional[str] = None) -> bool:
    if paragraph is None:
        return False
    text = paragraph_text_clean(paragraph)
    if not text:
        return False
    match = CAPTION_RE.match(text)
    if not match:
        return False
    return kind is None or match.group(1).lower() == kind.lower()


def get_prev_paragraph_element(elm) -> Optional[Paragraph]:
    prev = elm.getprevious()
    while prev is not None:
        if prev.tag == qn("w:p"):
            paragraph = Paragraph(prev, elm.getparent())
            if is_empty_paragraph(paragraph):
                prev = prev.getprevious()
                continue
            return paragraph
        if prev.tag == qn("w:tbl"):
            return None
        prev = prev.getprevious()
    return None


def get_next_paragraph_element(elm) -> Optional[Paragraph]:
    nxt = elm.getnext()
    while nxt is not None:
        if nxt.tag == qn("w:p"):
            paragraph = Paragraph(nxt, elm.getparent())
            if is_empty_paragraph(paragraph):
                nxt = nxt.getnext()
                continue
            return paragraph
        if nxt.tag == qn("w:tbl"):
            return None
        nxt = nxt.getnext()
    return None


def insert_paragraph_before_element(elm, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    elm.addprevious(new_p)
    paragraph = Paragraph(new_p, elm.getparent())
    if text:
        paragraph.add_run(text)
    return paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_run_font(run, name: str, size_pt: int, *, bold: Optional[bool] = None,
                 italic: Optional[bool] = None, underline: Optional[bool] = None,
                 all_caps: Optional[bool] = None) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if underline is not None:
        run.font.underline = underline
    if all_caps is not None:
        run.font.all_caps = all_caps

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:eastAsia"), name)


def sanitize_run_formatting(run) -> None:
    run.font.highlight_color = None
    run.font.subscript = False
    run.font.superscript = False
    run.font.strike = False
    run.font.double_strike = False
    run.font.shadow = False
    run.font.emboss = False
    run.font.imprint = False
    run.font.outline = False
    run.font.small_caps = False

    rpr = run._element.get_or_add_rPr()
    for tag in (
        "w:highlight", "w:vertAlign", "w:em", "w:effect", "w:fitText",
        "w:specVanish", "w:webHidden", "w:strike", "w:dstrike",
        "w:shadow", "w:outline", "w:emboss", "w:imprint", "w:smallCaps",
    ):
        elem = rpr.find(qn(tag))
        if elem is not None:
            rpr.remove(elem)

    for tag in ("w:spacing", "w:position", "w:kern", "w:w"):
        elem = rpr.find(qn(tag))
        if elem is not None:
            rpr.remove(elem)


def ensure_paragraph_style(doc: Document, style_name: str) -> None:
    try:
        doc.styles[style_name]
    except KeyError:
        doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)


def configure_document_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.font.color.rgb = BLACK
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(key), "Times New Roman")

    ensure_paragraph_style(doc, "FigureCaption")
    ensure_paragraph_style(doc, "TableCaption")


def set_page_layout(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)
        section.header_distance = Cm(0.9)


def set_auto_hyphenation(doc: Document, enabled: bool) -> None:
    settings = doc.settings.element
    auto = settings.find(qn("w:autoHyphenation"))
    if auto is None:
        auto = OxmlElement("w:autoHyphenation")
        settings.append(auto)
    auto.set(qn("w:val"), "true" if enabled else "false")


def strip_soft_hyphens(text: str) -> str:
    return text.replace(SOFT_HYPHEN, "")


def sanitize_runs_in_paragraph(paragraph: Paragraph, font_name: str, font_size: int,
                               *, bold: bool = False, italic: bool = False,
                               underline: bool = False, all_caps: bool = False) -> None:
    for run in paragraph.runs:
        if run_has_drawing(run):
            continue
        run.text = strip_soft_hyphens(run.text)
        sanitize_run_formatting(run)
        set_run_font(run, font_name, font_size, bold=bold, italic=italic,
                     underline=underline, all_caps=all_caps)


def format_body_paragraph(paragraph: Paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.widow_control = True
    pf.keep_together = False
    pf.keep_with_next = False
    sanitize_runs_in_paragraph(paragraph, "Times New Roman", 14)


def heading_base_text(text: str) -> str:
    text = strip_soft_hyphens(text).strip()
    text = LEADING_NUMBER_RE.sub("", text)
    return text.rstrip(".").strip().upper()


def format_heading(paragraph: Paragraph, number: int) -> None:
    title = heading_base_text(paragraph.text)
    clear_paragraph(paragraph)
    run = paragraph.add_run(f"{number} {title}" if title else f"{number}")
    set_run_font(run, "Arial", 14, bold=True, italic=False, underline=False, all_caps=True)
    sanitize_run_formatting(run)

    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(1.25)
    pf.right_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    pf.keep_together = True
    pf.page_break_before = True
    pf.widow_control = True


def add_page_number_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run, "Times New Roman", 12, bold=False, italic=False, underline=False)
    sanitize_run_formatting(run)


def clear_header_footer_content(container) -> Paragraph:
    while len(container.paragraphs) > 1:
        p = container.paragraphs[-1]._element
        p.getparent().remove(p)
    if container.paragraphs:
        p = container.paragraphs[0]
        clear_paragraph(p)
        return p
    return container.add_paragraph()


def set_headers_footers(doc: Document, meta: Meta) -> None:
    header_text = f"{meta.fio}, {meta.group}, Отчет по лабораторной работе {meta.lab_number}"

    for section in doc.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        header_p = clear_header_footer_content(section.header)
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_pf = header_p.paragraph_format
        header_pf.space_before = Pt(0)
        header_pf.space_after = Pt(0)
        header_pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        header_pf.widow_control = True
        header_run = header_p.add_run(header_text)
        set_run_font(header_run, "Times New Roman", 12, bold=False, italic=True, underline=False)
        sanitize_run_formatting(header_run)

        footer_p = clear_header_footer_content(section.footer)
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_pf = footer_p.paragraph_format
        footer_pf.space_before = Pt(0)
        footer_pf.space_after = Pt(0)
        footer_pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        footer_pf.widow_control = True
        add_page_number_field(footer_p)


def format_caption(paragraph: Paragraph, kind: str, number: int, title: str) -> None:
    title = title.strip() or PLACEHOLDER_TITLE
    clear_paragraph(paragraph)
    run = paragraph.add_run(f"{kind} {number} — {title}")
    sanitize_run_formatting(run)

    if kind == "Рисунок":
        set_run_font(run, "Times New Roman", 12, bold=False, italic=False, underline=False)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = paragraph.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0)
        pf.right_indent = Cm(0)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.keep_with_next = False
        pf.keep_together = True
        pf.widow_control = True
    else:
        set_run_font(run, "Times New Roman", 14, bold=False, italic=False, underline=False)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = paragraph.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(1.25)
        pf.right_indent = Cm(0)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_before = Pt(6)
        pf.space_after = Pt(0)
        pf.keep_with_next = True
        pf.keep_together = True
        pf.widow_control = True


def extract_caption_title(text: str) -> str:
    match = CAPTION_RE.match(text.strip())
    if not match:
        return PLACEHOLDER_TITLE
    title = match.group(3).strip()
    return title or PLACEHOLDER_TITLE


def ensure_table_caption(tbl: Table) -> Paragraph:
    prev_p = get_prev_paragraph_element(tbl._element)
    if is_caption_paragraph(prev_p, "Таблица"):
        return prev_p
    return insert_paragraph_before_element(tbl._element, f"Таблица 0 — {PLACEHOLDER_TITLE}")


def ensure_figure_caption(paragraph: Paragraph) -> Paragraph:
    next_p = get_next_paragraph_element(paragraph._element)
    if is_caption_paragraph(next_p, "Рисунок"):
        return next_p
    return insert_paragraph_after(paragraph, f"Рисунок 0 — {PLACEHOLDER_TITLE}")


def set_table_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def remove_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for elem in list(tr_pr):
        if elem.tag == qn("w:tblHeader"):
            tr_pr.remove(elem)


def cell_has_drawing(cell: _Cell) -> bool:
    return any(paragraph_has_drawing(p) for p in cell.paragraphs)


def estimate_cell_weight(cell: _Cell) -> float:
    text = " ".join((p.text or "").strip() for p in cell.paragraphs).strip()
    if cell_has_drawing(cell):
        # Картинкам и содержимому вокруг них нужна заметно большая ширина.
        return 10.0
    if not text:
        return 0.5

    compact = re.sub(r"\s+", " ", text)
    words = re.findall(r"\S+", compact)
    longest_word = max((len(w) for w in words), default=0)
    total_len = len(compact)
    line_count = max(1, len(cell.paragraphs))

    numeric_like = bool(re.fullmatch(r"[\d\s.,:;/%()+\-–—№A-Za-zА-Яа-я]+", compact)) and longest_word <= 14
    if numeric_like and total_len <= 20:
        return 1.5

    # Для текстовых столбцов ширину определяем по смеси общего объёма текста
    # и длины самого длинного слова, чтобы не было узких колонок с частыми переносами.
    weight = 1.0 + min(6.5, total_len / max(1.0, 18 - min(longest_word, 10) * 0.5))
    weight += min(3.5, longest_word * 0.18)
    weight += min(1.5, (line_count - 1) * 0.25)
    return max(1.5, min(13.0, weight))


def get_table_column_count(table: Table) -> int:
    return max((len(row.cells) for row in table.rows), default=0)


def get_table_usable_width_cm(table: Table, fallback: float = 16.8) -> float:
    parent = table._parent
    doc = parent if isinstance(parent, _Document) else None
    while doc is None and hasattr(parent, "_parent"):
        parent = parent._parent
        if isinstance(parent, _Document):
            doc = parent
            break
    if doc is None or not doc.sections:
        return fallback
    section = doc.sections[0]
    return max(8.0, (section.page_width - section.left_margin - section.right_margin) / 360000.0)


def adjust_table_column_widths(table: Table, usable_width_cm: Optional[float] = None) -> None:
    if not table.rows:
        return
    col_count = get_table_column_count(table)
    if col_count <= 1:
        return

    if usable_width_cm is None:
        usable_width_cm = get_table_usable_width_cm(table)

    desired: list[float] = []
    min_widths: list[float] = []
    max_widths: list[float] = []

    for col_idx in range(col_count):
        cell_weights: list[float] = []
        has_drawing = False
        short_count = 0
        nonempty_count = 0

        for row_idx, row in enumerate(table.rows):
            if col_idx >= len(row.cells):
                continue
            cell = row.cells[col_idx]
            w = estimate_cell_weight(cell)
            # Шапка влияет сильнее, потому что по ней обычно понятен смысл столбца.
            if row_idx == 0:
                w *= 1.15
            cell_weights.append(w)

            txt = re.sub(r"\s+", " ", " ".join((p.text or "").strip() for p in cell.paragraphs)).strip()
            if txt:
                nonempty_count += 1
                if len(txt) <= 12:
                    short_count += 1
            if cell_has_drawing(cell):
                has_drawing = True

        if not cell_weights:
            cell_weights = [1.5]

        cell_weights = sorted(cell_weights)
        max_w = cell_weights[-1]
        p75 = cell_weights[max(0, int(len(cell_weights) * 0.75) - 1)]
        avg = sum(cell_weights) / len(cell_weights)
        short_ratio = (short_count / nonempty_count) if nonempty_count else 0.0

        if has_drawing:
            desired_w = max(5.2, max_w * 0.7 + p75 * 0.5)
            mn, mx = 4.8, 9.5
        elif short_ratio >= 0.75:
            desired_w = max(1.7, min(3.2, avg * 0.9))
            mn, mx = 1.6, 3.6
        else:
            desired_w = max(2.4, max_w * 0.55 + p75 * 0.35 + avg * 0.1)
            mn, mx = 2.2, 7.8

        desired.append(desired_w)
        min_widths.append(mn)
        max_widths.append(mx)

    # Сначала ограничиваем разумными минимумами/максимумами, затем масштабируем в доступную ширину.
    widths_cm = [min(max(w, mn), mx) for w, mn, mx in zip(desired, min_widths, max_widths)]
    total = sum(widths_cm)
    if total <= 0:
        return

    scale = usable_width_cm / total
    widths_cm = [w * scale for w in widths_cm]

    # После масштабирования ещё раз прижимаем минимум/максимум и перераспределяем остаток.
    widths_cm = [min(max(w, mn), mx) for w, mn, mx in zip(widths_cm, min_widths, max_widths)]
    total = sum(widths_cm)
    diff = usable_width_cm - total
    if abs(diff) > 0.05:
        flexible = [i for i, (w, mn, mx) in enumerate(zip(widths_cm, min_widths, max_widths))
                    if (diff > 0 and w < mx - 0.05) or (diff < 0 and w > mn + 0.05)]
        if flexible:
            share = diff / len(flexible)
            for i in flexible:
                widths_cm[i] = min(max(widths_cm[i] + share, min_widths[i]), max_widths[i])

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    if tbl_pr is not None:
        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")

    for row in table.rows:
        for col_idx, cell in enumerate(row.cells):
            if col_idx >= len(widths_cm):
                continue
            width_cm = widths_cm[col_idx]
            cell.width = Cm(width_cm)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(Cm(width_cm).emu / 635)))


def format_table(table: Table, smart_widths: bool = False) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = not smart_widths

    if table.rows:
        remove_repeat_table_header(table.rows[0])
    if smart_widths:
        adjust_table_column_widths(table)

    for r_idx, row in enumerate(table.rows):
        set_table_row_cant_split(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.first_line_indent = Cm(0)
                pf.left_indent = Cm(0)
                pf.right_indent = Cm(0)
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                pf.keep_together = True
                pf.widow_control = True

                if paragraph_has_drawing(p):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    continue
                if is_caption_paragraph(p, "Рисунок"):
                    continue

                if r_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sanitize_runs_in_paragraph(p, "Times New Roman", 12, bold=True)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if not is_caption_paragraph(p):
                        sanitize_runs_in_paragraph(p, "Times New Roman", 12, bold=False)


def format_image_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.keep_with_next = True
    pf.keep_together = True
    pf.widow_control = True


def find_or_create_captions(doc: Document) -> None:
    blocks = list(iter_blocks_recursive(doc))
    for block in blocks:
        if isinstance(block, Table):
            ensure_table_caption(block)
        elif isinstance(block, Paragraph) and paragraph_has_drawing(block):
            ensure_figure_caption(block)


def renumber_and_format_captions(doc: Document, smart_widths: bool = False) -> None:
    table_no = 0
    figure_no = 0

    for block in iter_blocks_recursive(doc):
        if isinstance(block, Table):
            caption = ensure_table_caption(block)
            table_no += 1
            title = extract_caption_title(caption.text)
            format_caption(caption, "Таблица", table_no, title)
            format_table(block, smart_widths=smart_widths)
        elif isinstance(block, Paragraph) and paragraph_has_drawing(block):
            format_image_paragraph(block)
            caption = ensure_figure_caption(block)
            figure_no += 1
            title = extract_caption_title(caption.text)
            format_caption(caption, "Рисунок", figure_no, title)


def format_all_paragraphs(doc: Document) -> None:
    heading_no = 0
    for block in iter_block_items(doc):
        if isinstance(block, Table):
            continue
        paragraph = block

        if paragraph_has_drawing(paragraph):
            continue
        if is_caption_paragraph(paragraph):
            continue
        if is_empty_paragraph(paragraph):
            pf = paragraph.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            continue

        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name in HEADING_STYLE_NAMES:
            heading_no += 1
            format_heading(paragraph, heading_no)
        else:
            format_body_paragraph(paragraph)


def replace_text_wrapping_breaks_in_run(run) -> None:
    """Заменяет ручные переносы строки внутри run на пробелы, не трогая page/column breaks."""
    if run_has_drawing(run):
        return
    r = run._element
    children = list(r)
    for child in children:
        if child.tag != qn("w:br"):
            continue
        break_type = child.get(qn("w:type"))
        if break_type not in (None, "textWrapping", "line"):
            continue
        idx = list(r).index(child)
        text_elm = OxmlElement("w:t")
        text_elm.set(qn("xml:space"), "preserve")
        text_elm.text = " "
        r.remove(child)
        r.insert(idx, text_elm)


def remove_manual_line_breaks_in_text(doc: Document) -> None:
    for paragraph in iter_paragraphs_recursive(doc):
        for run in paragraph.runs:
            if run_has_drawing(run):
                continue
            replace_text_wrapping_breaks_in_run(run)
            text = run.text.replace("\v", " ").replace("\u2028", " ")
            run.text = strip_soft_hyphens(text)


def process_document(input_path: str, output_path: str, meta: Meta, smart_widths: bool = False) -> None:
    doc = Document(input_path)
    configure_document_defaults(doc)
    set_page_layout(doc)
    set_auto_hyphenation(doc, enabled=False)
    remove_manual_line_breaks_in_text(doc)
    set_headers_footers(doc, meta)
    format_all_paragraphs(doc)
    find_or_create_captions(doc)
    renumber_and_format_captions(doc, smart_widths=smart_widths)
    doc.save(output_path)


def build_output_path(input_path: str) -> str:
    root, ext = os.path.splitext(input_path)
    ext = ext or ".docx"
    return f"{root}_formatted{ext}"


def cli_main() -> int:
    parser = argparse.ArgumentParser(description="Форматирование .docx-отчета по шаблону лабораторной работы.")
    parser.add_argument("input", nargs="?", help="Входной .docx")
    parser.add_argument("output", nargs="?", help="Выходной .docx")
    parser.add_argument("--fio", help="ФИО для верхнего колонтитула")
    parser.add_argument("--group", help="Номер группы для верхнего колонтитула")
    parser.add_argument("--lab", help="Номер лабораторной работы")
    parser.add_argument("--smart-widths", action="store_true", help="Пытаться автоматически подгонять ширины столбцов таблиц")
    args = parser.parse_args()

    if not args.input:
        launch_gui()
        return 0

    if not args.output:
        args.output = build_output_path(args.input)

    missing = [name for name, value in (("--fio", args.fio), ("--group", args.group), ("--lab", args.lab)) if not value]
    if missing:
        parser.error("Для CLI-режима нужны параметры: --fio, --group, --lab")

    meta = Meta(fio=args.fio.strip(), group=args.group.strip(), lab_number=str(args.lab).strip())
    process_document(args.input, args.output, meta, smart_widths=args.smart_widths)
    print(f"Готово: {args.output}")
    return 0


def launch_gui() -> None:
    root = tk.Tk()
    root.title("Форматирование отчета .docx")
    root.geometry("720x430")
    root.resizable(False, False)

    input_var = tk.StringVar()
    output_var = tk.StringVar()
    fio_var = tk.StringVar()
    group_var = tk.StringVar()
    lab_var = tk.StringVar()
    smart_widths_var = tk.BooleanVar(value=False)

    def choose_input() -> None:
        path = filedialog.askopenfilename(
            title="Выберите .docx",
            filetypes=[("Word document", "*.docx")],
        )
        if path:
            input_var.set(path)
            if not output_var.get().strip():
                output_var.set(build_output_path(path))

    def choose_output() -> None:
        path = filedialog.asksaveasfilename(
            title="Сохранить как",
            defaultextension=".docx",
            filetypes=[("Word document", "*.docx")],
        )
        if path:
            output_var.set(path)

    def run_processing() -> None:
        input_path = input_var.get().strip()
        output_path = output_var.get().strip()
        fio = fio_var.get().strip()
        group = group_var.get().strip()
        lab = lab_var.get().strip()

        if not input_path:
            messagebox.showerror("Ошибка", "Выберите входной .docx файл.")
            return
        if not output_path:
            messagebox.showerror("Ошибка", "Укажите путь для сохранения результата.")
            return
        if not fio or not group or not lab:
            messagebox.showerror("Ошибка", "Заполните ФИО, группу и номер ЛР.")
            return

        try:
            process_document(
                input_path,
                output_path,
                Meta(fio=fio, group=group, lab_number=lab),
                smart_widths=smart_widths_var.get(),
            )
        except Exception as exc:  # noqa: BLE001
            messagebox.showerror("Ошибка", f"Не удалось обработать документ:\n{exc}")
            return

        messagebox.showinfo("Готово", f"Файл сохранен:\n{output_path}")

    pad_x = 12
    pad_y = 8

    row = 0
    tk.Label(root, text="Входной .docx").grid(row=row, column=0, sticky="w", padx=pad_x, pady=pad_y)
    tk.Entry(root, textvariable=input_var, width=64).grid(row=row, column=1, padx=pad_x, pady=pad_y)
    tk.Button(root, text="Выбрать", command=choose_input, width=12).grid(row=row, column=2, padx=pad_x, pady=pad_y)

    row += 1
    tk.Label(root, text="Выходной .docx").grid(row=row, column=0, sticky="w", padx=pad_x, pady=pad_y)
    tk.Entry(root, textvariable=output_var, width=64).grid(row=row, column=1, padx=pad_x, pady=pad_y)
    tk.Button(root, text="Сохранить как", command=choose_output, width=12).grid(row=row, column=2, padx=pad_x, pady=pad_y)

    row += 1
    tk.Label(root, text="ФИО").grid(row=row, column=0, sticky="w", padx=pad_x, pady=pad_y)
    tk.Entry(root, textvariable=fio_var, width=64).grid(row=row, column=1, padx=pad_x, pady=pad_y, columnspan=2, sticky="w")

    row += 1
    tk.Label(root, text="Номер группы").grid(row=row, column=0, sticky="w", padx=pad_x, pady=pad_y)
    tk.Entry(root, textvariable=group_var, width=32).grid(row=row, column=1, padx=pad_x, pady=pad_y, sticky="w")

    row += 1
    tk.Label(root, text="Номер ЛР").grid(row=row, column=0, sticky="w", padx=pad_x, pady=pad_y)
    tk.Entry(root, textvariable=lab_var, width=32).grid(row=row, column=1, padx=pad_x, pady=pad_y, sticky="w")

    row += 1
    tk.Checkbutton(
        root,
        text="Попробовать автоматически подгонять ширины столбцов таблиц",
        variable=smart_widths_var,
        onvalue=True,
        offvalue=False,
    ).grid(row=row, column=0, columnspan=3, sticky="w", padx=pad_x, pady=(10, 4))

    row += 1
    note = (
        "Автоматически добавляет подписи вида 'Таблица N — ...' и 'Рисунок N — ...'.\n"
        "Названия после тире можно потом дописать вручную.\n"
        "Умная подгонка ширин столбцов теперь старается сужать короткие столбцы и\n"
        "отдавать больше места длинным текстовым колонкам, но на сложной вёрстке\n"
        "её всё равно лучше включать выборочно."
    )
    tk.Label(root, text=note, justify="left", fg="#444444").grid(
        row=row, column=0, columnspan=3, sticky="w", padx=pad_x, pady=(16, 8)
    )

    row += 1
    page_url = "https://online.mospolytech.ru/mod/page/view.php?id=308998"
    link = tk.Label(
        root,
        text=page_url,
        fg="blue",
        cursor="hand2",
        justify="left",
        font=("Times New Roman", 11, "underline"),
    )
    link.grid(row=row, column=0, columnspan=3, sticky="w", padx=pad_x, pady=(4, 8))
    link.bind("<Button-1>", lambda _event: webbrowser.open_new(page_url))

    row += 1
    tk.Button(root, text="Отформатировать", command=run_processing, width=20, height=2).grid(
        row=row, column=0, columnspan=3, pady=20
    )

    root.mainloop()


if __name__ == "__main__":
    raise SystemExit(cli_main())
